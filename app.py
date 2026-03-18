"""
Text Summarizer - Streamlit Web Application
A comprehensive text summarization tool using both extractive and abstractive methods.
"""

import os
import re
# Set environment variable for Keras 3 compatibility BEFORE importing transformers
os.environ['TF_USE_LEGACY_KERAS'] = '1'

import streamlit as st
import time
from io import BytesIO

# Authentication imports
from auth import check_authentication, render_login_page, render_user_menu, get_current_user, logout

# Page configuration
st.set_page_config(
    page_title="Text Summarizer",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Theme styles (switchable from sidebar)
THEME_STYLES = {
    "Main Original Theme": "__CLASSIC_BLUE__",
    "Classic Blue": """
    :root { --accent-color: #1E88E5; }
    .main-header { font-size: 2.5rem; font-weight: bold; color: #1E88E5; text-align: center; margin-bottom: 1rem; }
    .sub-header { font-size: 1.2rem; color: #666; text-align: center; margin-bottom: 2rem; }
    .summary-box { background-color: #f0f7ff; padding: 20px; border-radius: 10px; border-left: 5px solid #1E88E5; }
    .stats-card { background-color: #f8f9fa; padding: 15px; border-radius: 8px; text-align: center; }
    .stButton>button { width: 100%; background-color: #1E88E5; color: white; }
    """,
    "Glass Morphism": """
    :root { --accent-color: #6366F1; }
    .stApp { background: linear-gradient(135deg, #f8fafc 0%, #eef2ff 55%, #f5f3ff 100%); }
    .main-header { font-size: 2.7rem; font-weight: 800; text-align: center; margin-bottom: 0.7rem; color: #4f46e5; letter-spacing: -0.5px; }
    .sub-header { font-size: 1.1rem; color: #64748b; text-align: center; margin-bottom: 1.5rem; }
    .summary-box { background: rgba(255,255,255,0.72); backdrop-filter: blur(10px); -webkit-backdrop-filter: blur(10px); padding: 20px; border-radius: 14px; border: 1px solid rgba(148,163,184,0.2); box-shadow: 0 8px 24px rgba(99,102,241,0.08); }
    .stats-card { background: rgba(255,255,255,0.78); backdrop-filter: blur(8px); border: 1px solid rgba(148,163,184,0.2); padding: 14px; border-radius: 12px; text-align: center; }
    .stButton>button { width: 100%; background: linear-gradient(135deg, #6366f1, #8b5cf6); color: white; border: none; border-radius: 10px; }
    """,
    "Dark Hacker": """
    :root { --accent-color: #00ff9c; }
    .stApp { background: radial-gradient(circle at top right, #1a1a1a 0%, #0f1115 55%, #0b0d10 100%); }
    .main-header { font-size: 2.6rem; font-weight: 800; color: #00ff9c; text-align: center; margin-bottom: 0.8rem; text-shadow: 0 0 8px rgba(0,255,156,0.35); }
    .sub-header { font-size: 1.1rem; color: #9efad8; text-align: center; margin-bottom: 1.6rem; }
    .summary-box { background-color: rgba(10,14,16,0.85); color: #d9ffe9; padding: 20px; border-radius: 10px; border-left: 4px solid #00ff9c; border: 1px solid rgba(0,255,156,0.2); }
    .stats-card { background-color: rgba(13,18,20,0.9); color: #d9ffe9; padding: 14px; border-radius: 8px; border: 1px solid rgba(0,255,156,0.2); text-align: center; }
    .stButton>button { width: 100%; background-color: #00c985; color: #06130d; border: none; border-radius: 8px; font-weight: 700; }
    [data-testid="stSidebar"] { background-color: #0d1214; }
    """,
    "Ocean Breeze": """
    :root { --accent-color: #0ea5e9; }
    .stApp { background: linear-gradient(140deg, #ecfeff 0%, #dbeafe 45%, #e0f2fe 100%); }
    .main-header { font-size: 2.6rem; font-weight: 800; color: #0369a1; text-align: center; margin-bottom: 0.8rem; }
    .sub-header { font-size: 1.1rem; color: #0f766e; text-align: center; margin-bottom: 1.5rem; }
    .summary-box { background-color: #f0f9ff; color: #0c4a6e; padding: 20px; border-radius: 12px; border-left: 4px solid #0ea5e9; border: 1px solid #bae6fd; }
    .stats-card { background-color: #ecfeff; color: #0c4a6e; padding: 14px; border-radius: 10px; border: 1px solid #a5f3fc; text-align: center; }
    .stButton>button { width: 100%; background: linear-gradient(135deg, #0ea5e9, #06b6d4); color: white; border: none; border-radius: 8px; }
    """,
    "Sunset Glow": """
    :root { --accent-color: #ea580c; }
    .stApp { background: linear-gradient(140deg, #fff7ed 0%, #ffedd5 52%, #fef3c7 100%); }
    .main-header { font-size: 2.6rem; font-weight: 800; color: #c2410c; text-align: center; margin-bottom: 0.8rem; }
    .sub-header { font-size: 1.1rem; color: #9a3412; text-align: center; margin-bottom: 1.5rem; }
    .summary-box { background-color: #fff7ed; color: #7c2d12; padding: 20px; border-radius: 12px; border-left: 4px solid #fb923c; border: 1px solid #fdba74; }
    .stats-card { background-color: #fffbeb; color: #7c2d12; padding: 14px; border-radius: 10px; border: 1px solid #fcd34d; text-align: center; }
    .stButton>button { width: 100%; background: linear-gradient(135deg, #f97316, #fb7185); color: white; border: none; border-radius: 8px; }
    """,
    "Forest Calm": """
    :root { --accent-color: #15803d; }
    .stApp { background: linear-gradient(145deg, #f0fdf4 0%, #dcfce7 58%, #ecfdf5 100%); }
    .main-header { font-size: 2.6rem; font-weight: 800; color: #166534; text-align: center; margin-bottom: 0.8rem; }
    .sub-header { font-size: 1.1rem; color: #14532d; text-align: center; margin-bottom: 1.5rem; }
    .summary-box { background-color: #f0fdf4; color: #14532d; padding: 20px; border-radius: 12px; border-left: 4px solid #22c55e; border: 1px solid #86efac; }
    .stats-card { background-color: #ecfdf5; color: #14532d; padding: 14px; border-radius: 10px; border: 1px solid #6ee7b7; text-align: center; }
    .stButton>button { width: 100%; background: linear-gradient(135deg, #16a34a, #059669); color: white; border: none; border-radius: 8px; }
    """,
    "Midnight Indigo": """
    :root { --accent-color: #818cf8; }
    .stApp { background: linear-gradient(150deg, #0f172a 0%, #1e1b4b 55%, #111827 100%); }
    .main-header { font-size: 2.6rem; font-weight: 800; color: #c7d2fe; text-align: center; margin-bottom: 0.8rem; }
    .sub-header { font-size: 1.1rem; color: #a5b4fc; text-align: center; margin-bottom: 1.5rem; }
    .summary-box { background-color: rgba(30,41,59,0.9); color: #e0e7ff; padding: 20px; border-radius: 12px; border-left: 4px solid #818cf8; border: 1px solid rgba(129,140,248,0.35); }
    .stats-card { background-color: rgba(31,41,55,0.9); color: #e0e7ff; padding: 14px; border-radius: 10px; border: 1px solid rgba(129,140,248,0.3); text-align: center; }
    .stButton>button { width: 100%; background: linear-gradient(135deg, #6366f1, #8b5cf6); color: white; border: none; border-radius: 8px; }
    [data-testid="stSidebar"] { background-color: #111827; }
    """,
    "Minimal Gray": """
    :root { --accent-color: #4b5563; }
    .stApp { background: linear-gradient(145deg, #f9fafb 0%, #f3f4f6 55%, #e5e7eb 100%); }
    .main-header { font-size: 2.5rem; font-weight: 750; color: #111827; text-align: center; margin-bottom: 0.9rem; }
    .sub-header { font-size: 1.1rem; color: #4b5563; text-align: center; margin-bottom: 1.5rem; }
    .summary-box { background-color: #ffffff; color: #111827; padding: 20px; border-radius: 10px; border-left: 4px solid #6b7280; border: 1px solid #d1d5db; }
    .stats-card { background-color: #f9fafb; color: #111827; padding: 14px; border-radius: 10px; border: 1px solid #d1d5db; text-align: center; }
    .stButton>button { width: 100%; background-color: #374151; color: white; border: none; border-radius: 8px; }
    """,
}


def apply_theme_css(theme_name: str):
    """Apply selected UI theme CSS."""
    css = THEME_STYLES.get(theme_name, THEME_STYLES["Classic Blue"])
    if css == "__CLASSIC_BLUE__":
        css = THEME_STYLES["Classic Blue"]
    st.markdown(f"<style>{css}</style>", unsafe_allow_html=True)


@st.cache_resource
def load_extractive_summarizer(method='textrank'):
    """Load and cache extractive summarizer."""
    from summarizers import ExtractiveSummarizer
    return ExtractiveSummarizer(method=method)


@st.cache_resource
def load_tfidf_summarizer():
    """Load and cache TF-IDF summarizer."""
    from summarizers import TFIDFSummarizer
    return TFIDFSummarizer()


@st.cache_resource
def load_advanced_summarizer():
    """Load and cache Advanced summarizer (uses all algorithms)."""
    from summarizers import AdvancedSummarizer
    return AdvancedSummarizer(use_clustering=True)


@st.cache_resource
def load_kmeans_summarizer():
    """Load and cache K-Means clustering summarizer."""
    from summarizers import KMeansClusterSummarizer
    return KMeansClusterSummarizer()


@st.cache_resource
def load_sbert_summarizer(model_name='all-MiniLM-L6-v2'):
    """Load and cache SBERT semantic summarizer."""
    from summarizers import SBERTSummarizer
    return SBERTSummarizer(model_name=model_name, use_mmr=True)


@st.cache_resource
def load_abstractive_summarizer(model_name):
    """Load and cache abstractive summarizer."""
    from summarizers import AbstractiveSummarizer
    return AbstractiveSummarizer(model_name=model_name)


def extract_text_from_upload(uploaded_file):
    """Extract text from uploaded file."""
    from utils import FileHandler
    handler = FileHandler()
    result = handler.extract_text(uploaded_file, filename=uploaded_file.name)
    return result


def detect_text_language(text):
    """Detect language of the input text."""
    from utils import detect_language
    return detect_language(text)


def resolve_active_language(input_text, auto_detect_lang, selected_language, language_options):
    """Resolve active language, persist to session state, and show status UI."""
    if input_text and auto_detect_lang:
        detected = detect_text_language(input_text)
        if detected['confidence'] > 0.5:
            st.info(f"🌐 Detected Language: {detected['flag']} {detected['name']} ({detected['confidence']:.0%} confidence)")
            active_language = detected['code']
        else:
            active_language = 'en'
    elif not auto_detect_lang:
        active_language = selected_language
        lang_info = language_options.get(selected_language, f"🌍 {selected_language}")
        st.info(f"🌐 Selected Language: {lang_info}")
    else:
        active_language = 'en'

    st.session_state['detected_language'] = active_language
    return active_language


def transcribe_audio(audio_data, language='en'):
    """Convert speech to text."""
    from utils import speech_to_text
    return speech_to_text(audio_data, language)


def generate_speech(text, language='en'):
    """Convert text to speech."""
    from utils import text_to_speech
    return text_to_speech(text, language)


def display_statistics(original_text, summary_text):
    """Display summary statistics."""
    original_words = len(original_text.split())
    summary_words = len(summary_text.split())
    compression = ((original_words - summary_words) / original_words * 100) if original_words > 0 else 0

    summary_characters = len(summary_text)
    summary_sentences = len([s for s in re.split(r'[.!?]+', summary_text.strip()) if s.strip()])
    unique_words = len({
        w.lower().strip(".,!?;:\"'()[]{}")
        for w in summary_text.split()
        if w.strip(".,!?;:\"'()[]{}")
    })

    # Additional detailed stats panel (keeps existing stats unchanged)
    with st.expander("📊 Summary Statistics", expanded=True):
        stat_col1, stat_col2, stat_col3, stat_col4 = st.columns(4)
        with stat_col1:
            st.metric("Characters", f"{summary_characters:,}")
        with stat_col2:
            st.metric("Words", f"{summary_words:,}")
        with stat_col3:
            st.metric("Sentences", f"{summary_sentences:,}")
        with stat_col4:
            st.metric("Unique Words", f"{unique_words:,}")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Original Words", f"{original_words:,}")
    
    with col2:
        st.metric("Summary Words", f"{summary_words:,}")
    
    with col3:
        st.metric("Compression", f"{compression:.1f}%")
    
    with col4:
        st.metric("Reduction", f"{original_words - summary_words:,} words")


def display_rouge_scores(summary_text, original_text):
    """Display ROUGE evaluation scores using original text as reference for content coverage."""
    if not original_text or not original_text.strip():
        st.warning("No original text to evaluate against.")
        return None
    
    if not summary_text or not summary_text.strip():
        st.warning("No summary text to evaluate.")
        return None
    
    try:
        from utils import SummaryEvaluator
        evaluator = SummaryEvaluator()
        scores = evaluator.evaluate(summary_text, original_text)
        
        st.markdown("### 📊 ROUGE Metrics (Content Coverage)")
        st.caption("Measures how well the summary captures the original text's content")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric(
                "ROUGE-1", 
                f"{scores['rouge1']['f1']:.1%}",
                help="Unigram overlap - word-level content coverage"
            )
        
        with col2:
            st.metric(
                "ROUGE-2", 
                f"{scores['rouge2']['f1']:.1%}",
                help="Bigram overlap - phrase-level content coverage"
            )
        
        with col3:
            st.metric(
                "ROUGE-L", 
                f"{scores['rougeL']['f1']:.1%}",
                help="Longest common subsequence"
            )
        
        with col4:
            st.metric(
                "ROUGE-Lsum", 
                f"{scores['rougeLsum']['f1']:.1%}",
                help="Summary-level ROUGE-L"
            )
        
        # Detailed scores expander
        with st.expander("📈 Detailed ROUGE Scores"):
            import pandas as pd
            data = {
                'Metric': ['ROUGE-1', 'ROUGE-2', 'ROUGE-L', 'ROUGE-Lsum'],
                'Precision': [
                    f"{scores['rouge1']['precision']:.2%}",
                    f"{scores['rouge2']['precision']:.2%}",
                    f"{scores['rougeL']['precision']:.2%}",
                    f"{scores['rougeLsum']['precision']:.2%}"
                ],
                'Recall': [
                    f"{scores['rouge1']['recall']:.2%}",
                    f"{scores['rouge2']['recall']:.2%}",
                    f"{scores['rougeL']['recall']:.2%}",
                    f"{scores['rougeLsum']['recall']:.2%}"
                ],
                'F1-Score': [
                    f"{scores['rouge1']['f1']:.2%}",
                    f"{scores['rouge2']['f1']:.2%}",
                    f"{scores['rougeL']['f1']:.2%}",
                    f"{scores['rougeLsum']['f1']:.2%}"
                ]
            }
            df = pd.DataFrame(data)
            st.table(df.set_index('Metric'))
            
            st.markdown("""
            **Interpretation:**
            - **Precision**: What % of summary words appear in original text
            - **Recall**: What % of original text words appear in summary
            - **F1-Score**: Balance between precision and recall
            """)
        
        return scores
        
    except ImportError as e:
        st.error(f"rouge-score not installed. Run: pip install rouge-score\nError: {e}")
        return None
    except Exception as e:
        st.error(f"Error calculating ROUGE scores: {e}")
        return None


def _estimate_syllables(word: str) -> int:
    """Estimate syllable count for readability scoring."""
    w = re.sub(r'[^a-z]', '', word.lower())
    if not w:
        return 1

    vowels = "aeiouy"
    syllables = 0
    prev_is_vowel = False

    for ch in w:
        is_vowel = ch in vowels
        if is_vowel and not prev_is_vowel:
            syllables += 1
        prev_is_vowel = is_vowel

    # Common adjustment for silent trailing 'e'
    if w.endswith('e') and syllables > 1:
        syllables -= 1

    return max(1, syllables)


def _flesch_reading_ease(text: str) -> float:
    """Compute approximate Flesch Reading Ease score (0-100+)."""
    words = re.findall(r"\b\w+\b", text)
    sentences = [s for s in re.split(r'[.!?]+', text) if s.strip()]

    if not words or not sentences:
        return 0.0

    word_count = len(words)
    sentence_count = max(1, len(sentences))
    syllable_count = sum(_estimate_syllables(w) for w in words)

    score = 206.835 - 1.015 * (word_count / sentence_count) - 84.6 * (syllable_count / word_count)
    return max(0.0, min(100.0, score))


def display_evaluation_dashboard(original_text, summary_text, method_label="Extractive"):
    """Display screenshot-style evaluation dashboard without touching model logic."""
    try:
        from utils import BasicMetrics

        if not original_text or not summary_text:
            return

        orig_words = max(1, len(original_text.split()))
        summ_words = max(1, len(summary_text.split()))

        compression_ratio = orig_words / summ_words
        reduction_pct = (1 - (summ_words / orig_words)) * 100
        coverage_pct = BasicMetrics.retention_ratio(original_text, summary_text) * 100
        novelty_pct = BasicMetrics.novelty_ratio(original_text, summary_text) * 100
        readability_score = _flesch_reading_ease(summary_text)
        avg_sentence_length = BasicMetrics.avg_sentence_length(summary_text)

        # Quality score heuristic from intrinsic metrics
        quality_score = (
            0.35 * coverage_pct +
            0.25 * max(0.0, min(100.0, reduction_pct)) +
            0.25 * readability_score +
            0.15 * (100 - novelty_pct)
        )
        quality_score = max(0.0, min(100.0, quality_score))

        st.markdown("### 📉 Evaluation Results")

        with st.expander(f"Evaluation: {method_label}", expanded=False):
            # Gauge chart (with fallback if plotly unavailable)
            try:
                import plotly.graph_objects as go

                fig = go.Figure(
                    go.Indicator(
                        mode="gauge+number",
                        value=quality_score,
                        number={"valueformat": ".1f"},
                        title={"text": "Quality Score", "font": {"size": 28}},
                        gauge={
                            "axis": {"range": [0, 100], "tickwidth": 1},
                            "bar": {"color": "#1f77b4"},
                            "steps": [
                                {"range": [0, 40], "color": "#f8cccc"},
                                {"range": [40, 70], "color": "#f3efc2"},
                                {"range": [70, 100], "color": "#c9f0c9"}
                            ],
                            "threshold": {
                                "line": {"color": "red", "width": 6},
                                "thickness": 0.75,
                                "value": 70
                            }
                        }
                    )
                )
                fig.update_layout(height=250, margin=dict(l=20, r=20, t=45, b=10))
                st.plotly_chart(fig, use_container_width=True)
            except Exception:
                st.metric("Quality Score", f"{quality_score:.1f}")
                st.progress(min(1.0, quality_score / 100.0))

            c1, c2, c3 = st.columns(3)

            with c1:
                st.markdown("#### 📊 Compression")
                st.markdown(
                    f"""
                    <div style="padding: 10px 12px; border: 1px solid #e5e7eb; border-radius: 10px; background: #fafafa;">
                        <div style="font-size: 0.85rem; color: #6b7280;">Compression Ratio</div>
                        <div style="font-size: 1.6rem; font-weight: 700; margin-bottom: 8px;">{compression_ratio:.2f}x</div>
                        <div style="font-size: 0.85rem; color: #6b7280;">Reduction</div>
                        <div style="font-size: 1.35rem; font-weight: 650;">{reduction_pct:.1f}%</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

            with c2:
                st.markdown("#### 📖 Coverage")
                st.markdown(
                    f"""
                    <div style="padding: 10px 12px; border: 1px solid #e5e7eb; border-radius: 10px; background: #fafafa;">
                        <div style="font-size: 0.85rem; color: #6b7280;">Word Coverage</div>
                        <div style="font-size: 1.6rem; font-weight: 700; margin-bottom: 8px;">{coverage_pct:.1f}%</div>
                        <div style="font-size: 0.85rem; color: #6b7280;">Novelty Ratio</div>
                        <div style="font-size: 1.35rem; font-weight: 650;">{novelty_pct:.1f}%</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

            with c3:
                st.markdown("#### ✏️ Readability")
                st.markdown(
                    f"""
                    <div style="padding: 10px 12px; border: 1px solid #e5e7eb; border-radius: 10px; background: #fafafa;">
                        <div style="font-size: 0.85rem; color: #6b7280;">Flesch Score</div>
                        <div style="font-size: 1.6rem; font-weight: 700; margin-bottom: 8px;">{readability_score:.1f}</div>
                        <div style="font-size: 0.85rem; color: #6b7280;">Avg Sentence Length</div>
                        <div style="font-size: 1.35rem; font-weight: 650;">{avg_sentence_length:.1f}</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

    except Exception as e:
        st.warning(f"Evaluation dashboard unavailable: {e}")


def main():
    """Main application function."""

    if 'ui_theme' not in st.session_state:
        st.session_state['ui_theme'] = "Classic Blue"

    # Apply selected theme before rendering the page body.
    apply_theme_css(st.session_state['ui_theme'])
    
    # Get current user
    user = get_current_user()
    
    # Header with welcome message
    col_title, col_user = st.columns([4, 1])
    with col_title:
        st.markdown('<div class="main-header">📝 Text Summarizer</div>', unsafe_allow_html=True)
        st.markdown('<div class="sub-header">Transform long texts into concise summaries using AI</div>', unsafe_allow_html=True)
    with col_user:
        if user:
            st.markdown(f"""
                <div style="text-align: right; padding: 10px;">
                    <span style="color: var(--accent-color, #1E88E5); font-weight: 600;">👤 {user.full_name or user.username}</span>
                </div>
            """, unsafe_allow_html=True)
            if st.button("🚪 Logout", key="header_logout"):
                logout()
    
    # Sidebar configuration
    with st.sidebar:
        st.header("⚙️ Settings")

        st.subheader("🎨 Appearance")
        st.selectbox(
            "Theme",
            options=list(THEME_STYLES.keys()),
            key="ui_theme",
            help="Choose a visual theme for the app"
        )

        st.divider()
        
        # Summarization method selection
        method = st.selectbox(
            "Summarization Method",
            ["Advanced (All Algorithms)", "Extractive (TF-IDF)", "Extractive (TextRank)", "K-Means Clustering", "Semantic (SBERT)", "Abstractive (AI)", "Compare All"],
            help="Choose how to generate the summary"
        )
        
        # AI Model selection - shown for Abstractive and Compare All
        if method in ["Abstractive (AI)", "Compare All"]:
            model_choice = st.selectbox(
                "AI Model",
                ["distilbart", "bart", "t5", "t5-base"],
                index=0,
                help="Select the AI model for summarization"
            )
        else:
            model_choice = "distilbart"  # Default value when not using AI
        
        st.divider()
        
        # Method-specific settings
        if method in ["Advanced (All Algorithms)", "Extractive (TF-IDF)", "Extractive (TextRank)", "K-Means Clustering", "Semantic (SBERT)", "Compare All"]:
            st.subheader("📝 Extractive Settings")
            num_sentences = st.number_input(
                "Number of Sentences",
                min_value=1,
                max_value=500,
                value=3,
                step=1,
                help="Number of key sentences to extract (1-50)"
            )
            
            if method == "Advanced (All Algorithms)":
                st.success("🚀 **Advanced** uses ALL algorithms: TF-IDF, TextRank, Position Weighting, Keyword Boosting, K-Means Clustering, Redundancy Removal")
            elif method == "Extractive (TF-IDF)":
                st.info("📊 **TF-IDF** scores sentences based on term importance across the document.")
            elif method == "Extractive (TextRank)":
                st.info("🔗 **TextRank** uses graph-based ranking based on sentence similarity.")
            elif method == "K-Means Clustering":
                st.info("🎯 **K-Means** clusters sentences by topic and selects representatives for diversity.")
            elif method == "Semantic (SBERT)":
                st.info("🧠 **SBERT** uses deep learning to understand sentence meaning for better similarity.")
        
        if method in ["Abstractive (AI)", "Compare All"]:
            st.subheader("🤖 Abstractive Settings")
            
            st.caption("Set custom summary length (in tokens)")
            col_min, col_max = st.columns(2)
            
            with col_min:
                min_length = st.number_input(
                    "Min Length",
                    min_value=5,
                    max_value=500,
                    value=30,
                    step=5,
                    help="Minimum tokens in summary"
                )
            
            with col_max:
                max_length = st.number_input(
                    "Max Length",
                    min_value=10,
                    max_value=1000,
                    value=150,
                    step=10,
                    help="Maximum tokens in summary"
                )
            
            # Validate min < max
            if min_length >= max_length:
                st.warning("⚠️ Min length should be less than Max length")
                max_length = min_length + 20
        else:
            # Default values when Abstractive settings are not shown
            max_length = 150
            min_length = 30
        
        # Ensure num_sentences has a default
        if method == "Abstractive (AI)":
            num_sentences = 3  # Default for abstractive-only
        
        st.divider()
        
        # Language Settings
        st.subheader("🌐 Language Settings")
        
        language_options = {
            'en': '🇬🇧 English',
            'hi': '🇮🇳 Hindi',
            'es': '🇪🇸 Spanish',
            'fr': '🇫🇷 French',
            'de': '🇩🇪 German'
        }
        
        auto_detect_lang = st.checkbox("Auto-detect language", value=True, help="Automatically detect the language of input text")
        
        if not auto_detect_lang:
            selected_language = st.selectbox(
                "Select Language",
                options=list(language_options.keys()),
                format_func=lambda x: language_options[x],
                help="Choose the language of your input text"
            )
        else:
            selected_language = 'auto'
        
        st.divider()
        
        # Voice Settings
        st.subheader("🎙️ Voice Features")
        
        enable_voice_input = st.checkbox("Enable Voice Input", value=False, help="Speak text to summarize")
        enable_tts = st.checkbox("Enable Text-to-Speech", value=True, help="Listen to the summary")
        
        st.divider()
        
        # ROUGE is always enabled - uses original text as reference
        enable_rouge = True  # Always enabled
        reference_summary = ""  # Will use original text automatically
        
        st.divider()
        
        # About section
        with st.expander("ℹ️ About - Algorithms Used"):
            st.markdown("""
            **Text Summarizer** implements these algorithms:
            
            **1. Preprocessing:**
            - Sentence Segmentation
            - Tokenization
            - Stop-word Removal
            - Lemmatization
            
            **2. Extractive Methods:**
            - **TF-IDF**: Term Frequency-Inverse Document Frequency
            - **TextRank**: Graph-based ranking (PageRank)
            - **Cosine Similarity**: Sentence comparison
            - **Position Weighting**: First/last sentence boost
            - **Length Filtering**: Quality control
            - **Keyword Boosting**: Domain relevance
            - **Redundancy Removal**: Duplicate filtering
            
            **3. Abstractive Method:**
            - **Transformers**: AI text generation (BART, T5)
            
            **Supported Formats**: TXT, PDF, DOCX
            """)
    
    # Initialize variables before tabs
    detected_language = 'en'  # Default language
    input_text = ""
    
    # Main content area - conditionally show voice tab
    if enable_voice_input:
        tab1, tab2, tab3 = st.tabs(["📝 Text Input", "📁 File Upload", "🎙️ Voice Input"])
    else:
        tab1, tab2 = st.tabs(["📝 Text Input", "📁 File Upload"])
    
    # Tab 1: Text Input
    with tab1:
        input_text = st.text_area(
            "Enter or paste your text here:",
            height=300,
            placeholder="Paste the text you want to summarize here...",
            help="Enter the text you want to summarize"
        )
        
        # Language detection display
        detected_language = resolve_active_language(
            input_text,
            auto_detect_lang,
            selected_language,
            language_options,
        )
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            summarize_text_btn = st.button("🚀 Summarize Text", use_container_width=True)
    
    # Tab 2: File Upload
    with tab2:
        uploaded_file = st.file_uploader(
            "Upload a document",
            type=['txt', 'pdf', 'docx'],
            help="Supported formats: TXT, PDF, DOCX"
        )
        
        if uploaded_file:
            st.success(f"✅ File uploaded: {uploaded_file.name}")
            
            # Extract text from file
            with st.spinner("Extracting text from file..."):
                result = extract_text_from_upload(uploaded_file)
            
            if result["success"]:
                input_text = result["text"]
                st.info(f"📊 Extracted {result['word_count']:,} words from {result['format'].upper()} file")
                
                # Language detection for uploaded file
                detected_language = resolve_active_language(
                    input_text,
                    auto_detect_lang,
                    selected_language,
                    language_options,
                )
                
                with st.expander("Preview extracted text"):
                    st.text_area("Extracted Text", input_text[:2000] + "..." if len(input_text) > 2000 else input_text, height=200, disabled=True)
                
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    summarize_file_btn = st.button("🚀 Summarize Document", use_container_width=True)
            else:
                st.error(f"❌ Error: {result['error']}")
                summarize_file_btn = False
        else:
            summarize_file_btn = False
            input_text = input_text if 'input_text' in dir() else ""
    
    # Tab 3: Voice Input (only if enabled)
    summarize_voice_btn = False
    voice_text = ""
    
    if enable_voice_input:
        with tab3:
            st.markdown("""
            **Speak to summarize!**
            
            Record your voice and the app will transcribe it for summarization.
            """)
            
            # Language selection for voice recognition
            voice_lang_options = {
                'en': '🇬🇧 English',
                'hi': '🇮🇳 Hindi',
                'es': '🇪🇸 Spanish',
                'fr': '🇫🇷 French',
                'de': '🇩🇪 German'
            }
            
            voice_language = st.selectbox(
                "Voice Language",
                options=list(voice_lang_options.keys()),
                format_func=lambda x: voice_lang_options[x],
                key="voice_lang_select",
                help="Select the language you will speak"
            )
            
            # Audio input (Streamlit 1.33+)
            st.info("🎤 Click the microphone to record your voice")
            audio_value = st.audio_input("Record audio", key="voice_recorder")
            
            if audio_value:
                st.audio(audio_value, format="audio/wav")
                
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    transcribe_btn = st.button("📝 Transcribe Audio", use_container_width=True)
                
                if transcribe_btn:
                    with st.spinner("Transcribing audio..."):
                        audio_bytes = audio_value.read()
                        result = transcribe_audio(audio_bytes, voice_language)
                    
                    if result['success']:
                        voice_text = result['text']
                        st.session_state['voice_transcribed_text'] = voice_text
                        st.success("✅ Audio transcribed successfully!")
                    else:
                        st.error(f"❌ {result['error']}")
                        st.session_state['voice_transcribed_text'] = ""
            
            # Show transcribed text if available
            if st.session_state.get('voice_transcribed_text'):
                voice_text = st.session_state['voice_transcribed_text']
                
                st.markdown("---")
                st.subheader("📝 Transcribed Text")
                
                # Editable text area
                edited_voice_text = st.text_area(
                    "Edit transcription if needed:",
                    value=voice_text,
                    height=200,
                    key="voice_text_edit"
                )
                
                st.info(f"📊 {len(edited_voice_text.split()):,} words")
                
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    summarize_voice_btn = st.button("🚀 Summarize Transcription", use_container_width=True)
                
                if summarize_voice_btn:
                    input_text = edited_voice_text
                    detected_language = voice_language
                    st.session_state['detected_language'] = voice_language
                
                # Clear button
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    if st.button("🗑️ Clear Transcription", use_container_width=True):
                        st.session_state['voice_transcribed_text'] = ""
                        st.rerun()
    
    # Process summarization
    should_summarize = (summarize_text_btn and input_text) or (summarize_file_btn and input_text) or (summarize_voice_btn and input_text)
    
    if should_summarize:
        if not input_text.strip():
            st.warning("⚠️ Please enter some text to summarize.")
        else:
            st.divider()
            st.header("📋 Summary Results")
            
            # Initialize result variables
            advanced_result = None
            tfidf_result = None
            textrank_result = None
            abstractive_result = None
            advanced_time = 0
            tfidf_time = 0
            textrank_time = 0
            abstractive_time = 0
            
            # Advanced Summarization (uses ALL algorithms)
            if method == "Advanced (All Algorithms)":
                with st.spinner("Generating summary using all algorithms..."):
                    start_time = time.time()
                    summarizer = load_advanced_summarizer()
                    advanced_result = summarizer.summarize(input_text, num_sentences=num_sentences)
                    advanced_time = time.time() - start_time
                
                st.subheader("📄 Summary (Advanced - All Algorithms)")
                st.markdown(f'<div class="summary-box">{advanced_result["summary"]}</div>', unsafe_allow_html=True)
                
                with st.expander("📊 Statistics"):
                    display_statistics(input_text, advanced_result["summary"])
                    st.caption(f"⏱️ Generated in {advanced_time:.2f} seconds")
                
                st.divider()
                display_rouge_scores(advanced_result["summary"], input_text)
                display_evaluation_dashboard(input_text, advanced_result["summary"], "Extractive")
                
                with st.expander("🔧 Algorithms Used"):
                    algorithms = advanced_result.get("algorithms_used", [])
                    cols = st.columns(3)
                    for i, algo in enumerate(algorithms):
                        if algo != "none":
                            cols[i % 3].markdown(f"✅ {algo.replace('_', ' ').title()}")
                    
                    keywords = advanced_result.get("keywords", [])
                    if keywords:
                        st.markdown("**Extracted Keywords:**")
                        st.write(", ".join(keywords[:10]))
            
            # TF-IDF Summarization
            if method in ["Extractive (TF-IDF)", "Compare All"]:
                with st.spinner("Generating TF-IDF summary..."):
                    start_time = time.time()
                    summarizer = load_tfidf_summarizer()
                    tfidf_result = summarizer.summarize(input_text, num_sentences=num_sentences)
                    tfidf_time = time.time() - start_time
                
                if method == "Extractive (TF-IDF)":
                    st.subheader("📄 Summary (TF-IDF)")
                    st.markdown(f'<div class="summary-box">{tfidf_result["summary"]}</div>', unsafe_allow_html=True)
                    
                    with st.expander("📊 Statistics"):
                        display_statistics(input_text, tfidf_result["summary"])
                        st.caption(f"⏱️ Generated in {tfidf_time:.2f} seconds | Method: TF-IDF")
                    
                    st.divider()
                    display_rouge_scores(tfidf_result["summary"], input_text)
                    display_evaluation_dashboard(input_text, tfidf_result["summary"], "Extractive")
            
            # TextRank Summarization
            if method in ["Extractive (TextRank)", "Compare All"]:
                with st.spinner("Generating TextRank summary..."):
                    start_time = time.time()
                    summarizer = load_extractive_summarizer('textrank')
                    textrank_result = summarizer.summarize(input_text, num_sentences=num_sentences)
                    textrank_time = time.time() - start_time
                
                if method == "Extractive (TextRank)":
                    st.subheader("📄 Summary (TextRank)")
                    st.markdown(f'<div class="summary-box">{textrank_result["summary"]}</div>', unsafe_allow_html=True)
                    
                    with st.expander("📊 Statistics"):
                        display_statistics(input_text, textrank_result["summary"])
                        st.caption(f"⏱️ Generated in {textrank_time:.2f} seconds | Method: TextRank")
                    
                    st.divider()
                    display_rouge_scores(textrank_result["summary"], input_text)
                    display_evaluation_dashboard(input_text, textrank_result["summary"], "Extractive")
            
            # K-Means Clustering Summarization
            kmeans_result = None
            kmeans_time = 0
            if method in ["K-Means Clustering", "Compare All"]:
                with st.spinner("Generating K-Means clustering summary..."):
                    start_time = time.time()
                    summarizer = load_kmeans_summarizer()
                    kmeans_summary = summarizer.summarize(input_text, num_sentences=num_sentences)
                    kmeans_time = time.time() - start_time
                    kmeans_result = {"summary": kmeans_summary}
                
                if method == "K-Means Clustering":
                    st.subheader("📄 Summary (K-Means Clustering)")
                    st.markdown(f'<div class="summary-box">{kmeans_result["summary"]}</div>', unsafe_allow_html=True)
                    
                    with st.expander("📊 Statistics"):
                        display_statistics(input_text, kmeans_result["summary"])
                        st.caption(f"⏱️ Generated in {kmeans_time:.2f} seconds | Method: K-Means Clustering")
                    
                    st.divider()
                    display_rouge_scores(kmeans_result["summary"], input_text)
                    display_evaluation_dashboard(input_text, kmeans_result["summary"], "Extractive")
                    
                    with st.expander("🎯 Cluster Information"):
                        cluster_info = summarizer.get_cluster_info(input_text)
                        for cluster_name, info in cluster_info.items():
                            if isinstance(info, dict) and 'top_terms' in info:
                                st.markdown(f"**{cluster_name.replace('_', ' ').title()}** ({info['size']} sentences)")
                                st.write(f"Topics: {', '.join(info['top_terms'])}")
            
            # SBERT Semantic Summarization
            sbert_result = None
            sbert_time = 0
            if method in ["Semantic (SBERT)", "Compare All"]:
                with st.spinner("Generating semantic summary using SBERT..."):
                    start_time = time.time()
                    summarizer = load_sbert_summarizer()
                    sbert_summary = summarizer.summarize(input_text, num_sentences=num_sentences)
                    sbert_time = time.time() - start_time
                    sbert_result = {"summary": sbert_summary}
                
                if method == "Semantic (SBERT)":
                    st.subheader("📄 Summary (SBERT Semantic)")
                    st.markdown(f'<div class="summary-box">{sbert_result["summary"]}</div>', unsafe_allow_html=True)
                    
                    with st.expander("📊 Statistics"):
                        display_statistics(input_text, sbert_result["summary"])
                        st.caption(f"⏱️ Generated in {sbert_time:.2f} seconds | Method: Sentence-BERT (MMR)")
                    
                    st.divider()
                    display_rouge_scores(sbert_result["summary"], input_text)
                    display_evaluation_dashboard(input_text, sbert_result["summary"], "Extractive")
                    
                    with st.expander("🧠 SBERT Details"):
                        st.write("**Model:** all-MiniLM-L6-v2")
                        st.write("**Selection Method:** Maximal Marginal Relevance (MMR)")
                        st.write("MMR balances relevance with diversity to avoid redundancy.")
            
            # Abstractive summarization
            if method in ["Abstractive (AI)", "Compare All"]:
                with st.spinner(f"Generating abstractive summary using {model_choice}..."):
                    start_time = time.time()
                    summarizer = load_abstractive_summarizer(model_choice)
                    abstractive_result = summarizer.summarize(
                        input_text,
                        max_length=max_length,
                        min_length=min_length
                    )
                    abstractive_time = time.time() - start_time
                
                if method == "Abstractive (AI)":
                    st.subheader("📄 Summary (AI Generated)")
                    st.markdown(f'<div class="summary-box">{abstractive_result["summary"]}</div>', unsafe_allow_html=True)
                    
                    with st.expander("📊 Statistics"):
                        display_statistics(input_text, abstractive_result["summary"])
                        st.caption(f"⏱️ Generated in {abstractive_time:.2f} seconds | Model: {model_choice}")
                    
                    st.divider()
                    display_rouge_scores(abstractive_result["summary"], input_text)
                    display_evaluation_dashboard(input_text, abstractive_result["summary"], "Abstractive")
            
            # Compare All view
            if method == "Compare All":
                st.subheader("🔍 Comparison of All Methods")
                
                # Display all summaries - Row 1
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.markdown("**📊 TF-IDF**")
                    st.markdown(f'<div class="summary-box" style="font-size: 0.9em;">{tfidf_result["summary"]}</div>', unsafe_allow_html=True)
                
                with col2:
                    st.markdown("**🔗 TextRank**")
                    st.markdown(f'<div class="summary-box" style="font-size: 0.9em;">{textrank_result["summary"]}</div>', unsafe_allow_html=True)
                
                with col3:
                    st.markdown("**🎯 K-Means**")
                    st.markdown(f'<div class="summary-box" style="font-size: 0.9em;">{kmeans_result["summary"]}</div>', unsafe_allow_html=True)
                
                # Row 2
                col4, col5 = st.columns(2)
                
                with col4:
                    st.markdown("**🧠 SBERT**")
                    st.markdown(f'<div class="summary-box" style="font-size: 0.9em;">{sbert_result["summary"]}</div>', unsafe_allow_html=True)
                
                with col5:
                    st.markdown("**🤖 Abstractive**")
                    st.markdown(f'<div class="summary-box" style="font-size: 0.9em;">{abstractive_result["summary"]}</div>', unsafe_allow_html=True)
                
                st.divider()
                st.subheader("📈 Performance Comparison")
                
                # Create comparison table
                comparison_data = {
                    "Metric": ["Words", "Time (s)", "Method Type"],
                    "TF-IDF": [
                        len(tfidf_result['summary'].split()),
                        f"{tfidf_time:.2f}",
                        "Extractive"
                    ],
                    "TextRank": [
                        len(textrank_result['summary'].split()),
                        f"{textrank_time:.2f}",
                        "Extractive"
                    ],
                    "K-Means": [
                        len(kmeans_result['summary'].split()),
                        f"{kmeans_time:.2f}",
                        "Clustering"
                    ],
                    "SBERT": [
                        len(sbert_result['summary'].split()),
                        f"{sbert_time:.2f}",
                        "Semantic"
                    ],
                    "Abstractive": [
                        len(abstractive_result['summary'].split()),
                        f"{abstractive_time:.2f}",
                        "Generative"
                    ]
                }
                
                import pandas as pd
                df = pd.DataFrame(comparison_data)
                st.table(df.set_index("Metric"))
                
                # ROUGE Comparison - Always shown using original text as reference
                st.divider()
                st.subheader("📊 ROUGE Score Comparison")
                
                try:
                    from utils import SummaryEvaluator
                    evaluator = SummaryEvaluator()
                    
                    # Calculate ROUGE for all methods using original input_text as reference
                    rouge_comparison = {
                        "Metric": ["ROUGE-1 F1", "ROUGE-2 F1", "ROUGE-L F1"],
                        "TF-IDF": [],
                        "TextRank": [],
                        "K-Means": [],
                        "SBERT": [],
                        "Abstractive": []
                    }
                    
                    method_summaries = {
                        "TF-IDF": tfidf_result['summary'],
                        "TextRank": textrank_result['summary'],
                        "K-Means": kmeans_result['summary'],
                        "SBERT": sbert_result['summary'],
                        "Abstractive": abstractive_result['summary']
                    }
                    
                    for method_name, summary in method_summaries.items():
                        scores = evaluator.evaluate(summary, input_text)
                        rouge_comparison[method_name] = [
                            f"{scores['rouge1']['f1']:.1%}",
                            f"{scores['rouge2']['f1']:.1%}",
                            f"{scores['rougeL']['f1']:.1%}"
                        ]
                    
                    rouge_df = pd.DataFrame(rouge_comparison)
                    st.table(rouge_df.set_index("Metric"))
                    
                    # Find best method
                    best_scores = {}
                    for method_name, summary in method_summaries.items():
                        scores = evaluator.evaluate(summary, input_text)
                        best_scores[method_name] = scores['rouge1']['f1']
                    
                    best_method = max(best_scores, key=best_scores.get)
                    st.success(f"🏆 **Best Method (ROUGE-1):** {best_method} with {best_scores[best_method]:.1%} F1 score")
                    
                except ImportError:
                    st.warning("Install rouge-score for ROUGE evaluation: pip install rouge-score")
            
            # Download option
            st.divider()
            
            if method == "Compare All":
                download_text = f"""TEXT SUMMARIZATION RESULTS
{'='*50}

ORIGINAL TEXT LENGTH: {len(input_text.split())} words

TF-IDF SUMMARY (Extractive):
{'-'*30}
{tfidf_result['summary']}

TEXTRANK SUMMARY (Extractive):
{'-'*30}
{textrank_result['summary']}

K-MEANS CLUSTERING SUMMARY:
{'-'*30}
{kmeans_result['summary']}

SBERT SEMANTIC SUMMARY:
{'-'*30}
{sbert_result['summary']}

ABSTRACTIVE SUMMARY (AI Generated - {model_choice}):
{'-'*30}
{abstractive_result['summary']}
"""
            elif method == "Advanced (All Algorithms)":
                download_text = f"""ADVANCED TEXT SUMMARIZATION RESULTS
{'='*50}

ORIGINAL TEXT LENGTH: {len(input_text.split())} words
ALGORITHMS USED: {', '.join(advanced_result.get('algorithms_used', []))}

SUMMARY:
{'-'*30}
{advanced_result['summary']}

EXTRACTED KEYWORDS: {', '.join(advanced_result.get('keywords', [])[:10])}
"""
            elif method == "Extractive (TF-IDF)":
                download_text = tfidf_result["summary"]
            elif method == "Extractive (TextRank)":
                download_text = textrank_result["summary"]
            elif method == "K-Means Clustering":
                download_text = kmeans_result["summary"]
            elif method == "Semantic (SBERT)":
                download_text = sbert_result["summary"]
            elif method == "Abstractive (AI)":
                download_text = abstractive_result["summary"]
            else:
                download_text = abstractive_result["summary"] if abstractive_result else ""
            
            # Download section - All three formats available at once
            st.subheader("📥 Download Summary")
            
            # Prepare all download formats
            # 1. Text format - ready as-is
            txt_data = download_text
            
            # 2. PDF format
            try:
                from fpdf import FPDF
                from io import BytesIO
                
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", "B", 16)
                pdf.cell(0, 10, "Text Summary", ln=True, align="C")
                pdf.ln(10)
                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 10, f"Method: {method}", ln=True)
                pdf.ln(5)
                pdf.set_font("Arial", "", 11)
                safe_text = download_text.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 7, safe_text)
                
                pdf_output = BytesIO()
                pdf.output(pdf_output)
                pdf_data = pdf_output.getvalue()
                pdf_ready = True
            except Exception as e:
                pdf_data = None
                pdf_ready = False
                pdf_error = str(e)
            
            # 3. DOCX format
            try:
                from docx import Document
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from io import BytesIO
                
                doc = Document()
                title = doc.add_heading("Text Summary", 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph(f"Method: {method}")
                doc.add_paragraph(f"Original Text Length: {len(input_text.split())} words")
                doc.add_paragraph("")
                doc.add_heading("Summary", level=1)
                doc.add_paragraph(download_text)
                
                docx_output = BytesIO()
                doc.save(docx_output)
                docx_output.seek(0)
                docx_data = docx_output.getvalue()
                docx_ready = True
            except Exception as e:
                docx_data = None
                docx_ready = False
                docx_error = str(e)
            
            # Display all three download buttons
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.download_button(
                    label="📄 Download TXT",
                    data=txt_data,
                    file_name="summary.txt",
                    mime="text/plain",
                    use_container_width=True,
                    key="download_txt_btn"
                )
            
            with col2:
                if pdf_ready:
                    st.download_button(
                        label="📕 Download PDF",
                        data=pdf_data,
                        file_name="summary.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                        key="download_pdf_btn"
                    )
                else:
                    st.error(f"PDF error: {pdf_error}")
            
            with col3:
                if docx_ready:
                    st.download_button(
                        label="📘 Download DOCX",
                        data=docx_data,
                        file_name="summary.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="download_docx_btn"
                    )
                else:
                    st.error(f"DOCX error: {docx_error}")
            
            # Text-to-Speech Section
            if enable_tts and download_text:
                st.divider()
                st.subheader("🔊 Listen to Summary")
                
                # Get language for TTS from session state
                tts_language = st.session_state.get('detected_language', 'en')
                
                # Language display
                tts_lang_names = {
                    'en': '🇬🇧 English',
                    'hi': '🇮🇳 Hindi',
                    'es': '🇪🇸 Spanish',
                    'fr': '🇫🇷 French',
                    'de': '🇩🇪 German'
                }
                
                # Allow user to change TTS language first so UI shows effective value
                tts_lang_override = st.selectbox(
                    "Change language",
                    options=['auto'] + list(tts_lang_names.keys()),
                    format_func=lambda x: "Auto-detect" if x == 'auto' else tts_lang_names.get(x, x),
                    key="tts_lang_override"
                )

                effective_tts_language = tts_language if tts_lang_override == 'auto' else tts_lang_override

                col1, col2 = st.columns([2, 1])
                with col1:
                    st.info(f"🌐 TTS Language: {tts_lang_names.get(effective_tts_language, 'English')}")
                with col2:
                    if tts_lang_override == 'auto':
                        st.caption("Using detected language")
                    else:
                        st.caption("Using manual override")
                
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    generate_audio_btn = st.button("🔊 Generate Audio", use_container_width=True, key="tts_btn")
                
                if generate_audio_btn:
                    with st.spinner("Generating audio..."):
                        tts_result = generate_speech(download_text, effective_tts_language)
                    
                    if tts_result['success']:
                        st.audio(tts_result['audio_data'], format='audio/mp3')
                        st.download_button(
                            label="📥 Download Audio (MP3)",
                            data=tts_result['audio_data'],
                            file_name="summary_audio.mp3",
                            mime="audio/mp3",
                            key="download_audio_btn"
                        )
                    else:
                        st.error(f"❌ {tts_result['error']}")


# Sample text for demo
def show_demo():
    """Show demo with sample text."""
    sample_text = """
    Artificial intelligence (AI) is intelligence demonstrated by machines, as opposed to natural intelligence 
    displayed by animals including humans. AI research has been defined as the field of study of intelligent agents, 
    which refers to any system that perceives its environment and takes actions that maximize its chance of achieving 
    its goals. The term "artificial intelligence" had previously been used to describe machines that mimic and display 
    "human" cognitive skills that are associated with the human mind, such as "learning" and "problem-solving". 
    This definition has since been rejected by major AI researchers who now describe AI in terms of rationality 
    and acting rationally, which does not limit how intelligence can be articulated.
    
    AI applications include advanced web search engines, recommendation systems, understanding human speech, 
    self-driving cars, generative or creative tools, automated decision-making, and competing at the highest level 
    in strategic game systems. As machines become increasingly capable, tasks considered to require "intelligence" 
    are often removed from the definition of AI, a phenomenon known as the AI effect. For instance, optical character 
    recognition is frequently excluded from things considered to be AI, having become a routine technology.
    
    Artificial intelligence was founded as an academic discipline in 1956, and in the years since it has experienced 
    several waves of optimism, followed by disappointment and the loss of funding, followed by new approaches, 
    success, and renewed funding. AI research has tried and discarded many different approaches, including 
    simulating the brain, modeling human problem solving, formal logic, large databases of knowledge, and imitating 
    animal behavior. In the first decades of the 21st century, highly mathematical and statistical machine learning 
    has dominated the field, and this technique has proved highly successful, helping to solve many challenging 
    problems throughout industry and academia.
    """
    return sample_text


if __name__ == "__main__":
    # Check if user is authenticated
    if not check_authentication():
        # Show vibrant login page
        render_login_page()
    else:
        # User is logged in, show main app
        main()
        # Add user menu to sidebar
        render_user_menu()
