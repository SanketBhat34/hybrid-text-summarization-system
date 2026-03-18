"""
File Handler Module
Handles reading and processing various file formats for text extraction.
Supports: PDF, TXT, DOCX
"""

import os
from typing import Union, BinaryIO
import io


class FileHandler:
    """
    Handles file uploads and text extraction from various formats.
    """
    
    SUPPORTED_EXTENSIONS = {'.txt', '.pdf', '.docx', '.doc'}
    
    def __init__(self):
        """Initialize the file handler."""
        pass
    
    @staticmethod
    def get_file_extension(filename: str) -> str:
        """Get file extension in lowercase."""
        return os.path.splitext(filename)[1].lower()
    
    @staticmethod
    def is_supported(filename: str) -> bool:
        """Check if file format is supported."""
        ext = FileHandler.get_file_extension(filename)
        return ext in FileHandler.SUPPORTED_EXTENSIONS
    
    def extract_text(self, file: Union[str, BinaryIO], filename: str = None) -> dict:
        """
        Extract text from a file.
        
        Args:
            file: File path (str) or file-like object
            filename: Original filename (required if file is file-like object)
            
        Returns:
            Dictionary with extracted text and metadata
        """
        if isinstance(file, str):
            filename = file
            if not os.path.exists(file):
                return {"success": False, "error": "File not found", "text": ""}
        
        ext = self.get_file_extension(filename)
        
        if ext == '.txt':
            return self._extract_from_txt(file)
        elif ext == '.pdf':
            return self._extract_from_pdf(file)
        elif ext in ['.docx', '.doc']:
            return self._extract_from_docx(file)
        else:
            return {
                "success": False, 
                "error": f"Unsupported format: {ext}",
                "text": ""
            }
    
    def _extract_from_txt(self, file: Union[str, BinaryIO]) -> dict:
        """Extract text from TXT file."""
        try:
            if isinstance(file, str):
                with open(file, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                # Handle file-like object
                content = file.read()
                if isinstance(content, bytes):
                    text = content.decode('utf-8')
                else:
                    text = content
            
            return {
                "success": True,
                "text": text,
                "format": "txt",
                "word_count": len(text.split()),
                "char_count": len(text)
            }
        except Exception as e:
            return {"success": False, "error": str(e), "text": ""}
    
    def _extract_from_pdf(self, file: Union[str, BinaryIO]) -> dict:
        """Extract text from PDF file."""
        try:
            import PyPDF2
            
            if isinstance(file, str):
                with open(file, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = self._extract_pdf_text(reader)
            else:
                # Handle file-like object
                reader = PyPDF2.PdfReader(file)
                text = self._extract_pdf_text(reader)
            
            return {
                "success": True,
                "text": text,
                "format": "pdf",
                "word_count": len(text.split()),
                "char_count": len(text)
            }
        except ImportError:
            return {
                "success": False, 
                "error": "PyPDF2 not installed. Run: pip install PyPDF2",
                "text": ""
            }
        except Exception as e:
            return {"success": False, "error": str(e), "text": ""}
    
    def _extract_pdf_text(self, reader) -> str:
        """Extract text from all pages of PDF."""
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return '\n'.join(text_parts)
    
    def _extract_from_docx(self, file: Union[str, BinaryIO]) -> dict:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            
            if isinstance(file, str):
                doc = Document(file)
            else:
                # Handle file-like object
                doc = Document(file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            text = '\n'.join(text_parts)
            
            return {
                "success": True,
                "text": text,
                "format": "docx",
                "word_count": len(text.split()),
                "char_count": len(text)
            }
        except ImportError:
            return {
                "success": False, 
                "error": "python-docx not installed. Run: pip install python-docx",
                "text": ""
            }
        except Exception as e:
            return {"success": False, "error": str(e), "text": ""}


def extract_text_from_file(file_path: str) -> str:
    """
    Convenience function to extract text from a file.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Extracted text or error message
    """
    handler = FileHandler()
    result = handler.extract_text(file_path)
    
    if result["success"]:
        return result["text"]
    else:
        raise ValueError(result["error"])


def get_supported_formats() -> list:
    """Return list of supported file formats."""
    return list(FileHandler.SUPPORTED_EXTENSIONS)
